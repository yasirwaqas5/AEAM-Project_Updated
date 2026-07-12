"""
aeam/tests/test_phase_b1_3_ingestion.py

Phase B1.3 — Ingestion Processing Pipeline.

Two layers, both self-contained (no Qdrant / Redis / Postgres needed):

1. Text extraction (aeam.ingestion.extraction) — every Tier 1+2 format is
   exercised against REAL synthetic files (a genuine DOCX/XLSX/PDF is generated
   in-test), plus the deferred-category and error paths.

2. Processor orchestration (aeam.ingestion.processor.DocumentIngestJobProcessor)
   driven through the real IngestionWorker against a temp SQLite DatabaseClient
   and a temp on-disk BlobStore, with a FAKE IngestionPipeline so the chunk/
   embed/index step is isolated from Qdrant. Verifies the full state machine,
   document lifecycle, version chunk_id capture, dedup no-op, and failure
   handling.
"""

from __future__ import annotations

import io

import pytest

from aeam.ingestion.extraction import (
    extract_text,
    can_extract,
    ExtractionError,
    UnsupportedCategoryError,
    PROCESSABLE_CATEGORIES,
    DEFERRED_CATEGORIES,
)
from aeam.ingestion.processor import DocumentIngestJobProcessor, ProcessingError
from aeam.ingestion.worker import IngestionWorker
from aeam.integrations.database import DatabaseClient
from aeam.storage.blob_store import LocalDiskBlobStore
from aeam.registry.models import (
    AssetStatus, Document, IngestionJob, JobStatus, JobType, ParentType,
    Source, SourceKind, Version,
)
from aeam.registry.repositories import (
    DocumentRepository, IngestionJobRepository, SourceRepository, VersionRepository,
)


# ===========================================================================
# 1. Extraction
# ===========================================================================

def test_extract_markdown():
    r = extract_text(b"# Title\n\nSome **markdown** body text.", "markdown", "a.md")
    assert "Title" in r.text and "body text" in r.text


def test_extract_log():
    r = extract_text(b"2026-07-12 ERROR payment failed\nline2", "log", "a.log")
    assert "payment failed" in r.text


def test_extract_json_pretty_prints():
    r = extract_text(b'{"service":"payments","errors":[1,2,3]}', "json", "a.json")
    assert "payments" in r.text and "\n" in r.text  # re-serialised with indent


def test_extract_json_invalid_raises():
    with pytest.raises(ExtractionError) as exc:
        extract_text(b"{not valid json", "json", "a.json")
    assert exc.value.reason == "invalid_json"


def test_extract_xml_recovers_text_and_unescapes():
    r = extract_text(
        b"<root><item>CPU spike</item><note>web-01 &amp; db</note></root>", "xml", "a.xml"
    )
    assert "CPU spike" in r.text and "web-01 & db" in r.text


def test_extract_csv():
    r = extract_text(b"name,value\nrevenue,100\ncost,40\n", "csv", "a.csv")
    assert "revenue" in r.text
    assert r.detail["rows"] == 2 and r.detail["columns"] == 2


def test_extract_excel_real_xlsx():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Metrics"
    ws.append(["kpi", "value"])
    ws.append(["latency_ms", 220])
    ws.append(["error_rate", 0.02])
    buf = io.BytesIO()
    wb.save(buf)
    r = extract_text(buf.getvalue(), "excel", "a.xlsx")
    assert "latency_ms" in r.text and "Sheet: Metrics" in r.text
    assert r.detail["rows"] == 2


def test_extract_docx_real_docx():
    import docx
    d = docx.Document()
    d.add_paragraph("Runbook: restart the payment service.")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Metric"
    table.rows[0].cells[1].text = "Threshold"
    buf = io.BytesIO()
    d.save(buf)
    r = extract_text(buf.getvalue(), "docx", "a.docx")
    assert "payment service" in r.text and "Metric | Threshold" in r.text


def _make_minimal_pdf(text: str) -> bytes:
    """A byte-offset-correct minimal PDF with a real text layer pypdf can read."""
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>",
    ]
    stream = b"BT /F1 24 Tf 72 700 Td (" + text.encode("latin-1") + b") Tj ET"
    objs.append(b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream")
    objs.append(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    out = b"%PDF-1.4\n"
    offsets = []
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += str(i).encode() + b" 0 obj\n" + body + b"\nendobj\n"
    xref_pos = len(out)
    n = len(objs) + 1
    out += b"xref\n0 " + str(n).encode() + b"\n0000000000 65535 f \n"
    for off in offsets:
        out += ("%010d 00000 n \n" % off).encode()
    out += b"trailer\n<< /Size " + str(n).encode() + b" /Root 1 0 R >>\nstartxref\n"
    out += str(xref_pos).encode() + b"\n%%EOF"
    return out


def test_extract_pdf_real_pdf():
    r = extract_text(_make_minimal_pdf("Hello AEAM PDF text"), "pdf", "a.pdf")
    assert "AEAM PDF" in r.text
    assert r.detail["pages"] == 1


@pytest.mark.parametrize("category", ["image", "audio", "video"])
def test_extract_deferred_categories_raise(category):
    assert not can_extract(category)
    with pytest.raises(UnsupportedCategoryError) as exc:
        extract_text(b"\x00\x01binarydata", category, f"a.{category}")
    assert exc.value.reason == "deferred_category"
    assert exc.value.category == category


def test_extract_empty_raises():
    with pytest.raises(ExtractionError) as exc:
        extract_text(b"", "markdown", "a.md")
    assert exc.value.reason == "empty_content"


def test_category_registry_is_stable():
    assert PROCESSABLE_CATEGORIES == {
        "markdown", "log", "json", "xml", "csv", "excel", "pdf", "docx"
    }
    assert DEFERRED_CATEGORIES == {"image", "audio", "video"}


# ===========================================================================
# 2. Processor orchestration
# ===========================================================================

class _FakePipeline:
    """Stand-in for the RAG IngestionPipeline — records calls, no Qdrant."""

    collection = "aeam_documents"

    def __init__(self, chunk_ids=None):
        self.calls: list[dict] = []
        self._chunk_ids = chunk_ids if chunk_ids is not None else ["pt-0", "pt-1", "pt-2"]

    def ingest_document(self, text: str, metadata: dict) -> dict:
        self.calls.append({"text": text, "metadata": metadata})
        return {
            "collection": self.collection,
            "chunks_total": len(self._chunk_ids),
            "chunks_upserted": len(self._chunk_ids),
            "chunk_ids": list(self._chunk_ids),
            "doc_type": metadata.get("doc_type"),
            "source": metadata.get("source"),
            "date": metadata.get("date"),
        }


@pytest.fixture()
def db(tmp_path):
    client = DatabaseClient(database_url=f"sqlite:///{(tmp_path / 'b13.db').as_posix()}")
    yield client
    client.dispose()


@pytest.fixture()
def blob_store(tmp_path):
    return LocalDiskBlobStore(tmp_path / "blobs")


def _seed_upload(db, blob_store, data: bytes, category: str, filename: str,
                 doc_status: str = AssetStatus.PENDING):
    """Mimic what the ingress API does at upload: blob + source + doc + version + job."""
    blob_ref = blob_store.put(data)
    source_repo = SourceRepository(db)
    doc_repo = DocumentRepository(db)
    version_repo = VersionRepository(db)
    job_repo = IngestionJobRepository(db)

    source_id = source_repo.create(Source(name="Manual Upload", kind=SourceKind.UPLOAD))
    doc_id = doc_repo.create(Document(
        title=filename, source_id=source_id, origin_path=filename, doc_type=category,
        content_hash=blob_ref.content_hash, status=doc_status,
    ))
    version_repo.create(Version(
        parent_type=ParentType.DOCUMENT, parent_id=doc_id, version=1,
        content_hash=blob_ref.content_hash, blob_ref=blob_ref.uri, is_active=True,
    ))
    job_id = job_repo.create(IngestionJob(
        job_type=JobType.INGEST, source_id=source_id,
        parent_type=ParentType.DOCUMENT, parent_id=doc_id,
        status=JobStatus.QUEUED, content_hash=blob_ref.content_hash,
    ))
    return job_id, doc_id


def _drain(db, blob_store, pipeline):
    """Run one worker cycle with the real processor over the queued job."""
    job_repo = IngestionJobRepository(db)
    processor = DocumentIngestJobProcessor(blob_store=blob_store, ingestion_pipeline=pipeline, db=db)
    worker = IngestionWorker(job_repo=job_repo, processor=processor, poll_interval=1)
    return worker.run_once()


def test_processor_happy_path_indexes_and_finalises(db, blob_store):
    pipeline = _FakePipeline(chunk_ids=["a", "b", "c", "d"])
    job_id, doc_id = _seed_upload(db, blob_store, b"# Runbook\n\nRestart the service.", "markdown", "run.md")

    claimed = _drain(db, blob_store, pipeline)
    assert claimed is True

    job = IngestionJobRepository(db).get(job_id)
    doc = DocumentRepository(db).get(doc_id)
    version = VersionRepository(db).get_active(ParentType.DOCUMENT, doc_id)

    # Job reached DONE via EXTRACTING/INDEXING.
    assert job.status == JobStatus.DONE
    assert job.progress == 100
    # Document finalised.
    assert doc.status == AssetStatus.INDEXED
    assert doc.chunk_count == 4
    # Version captured the Qdrant point ids.
    assert version.chunk_ids == ["a", "b", "c", "d"]
    # Pipeline was called once with extracted text + required metadata.
    assert len(pipeline.calls) == 1
    md = pipeline.calls[0]["metadata"]
    assert md["doc_type"] == "markdown" and md["doc_id"] == doc_id
    assert "Restart the service" in pipeline.calls[0]["text"]


def test_processor_dedup_noop_when_already_indexed(db, blob_store):
    pipeline = _FakePipeline()
    job_id, doc_id = _seed_upload(
        db, blob_store, b"already indexed content", "markdown", "x.md",
        doc_status=AssetStatus.INDEXED,
    )

    assert _drain(db, blob_store, pipeline) is True

    job = IngestionJobRepository(db).get(job_id)
    assert job.status == JobStatus.DONE
    # No re-embedding for content already indexed.
    assert pipeline.calls == []


def test_processor_deferred_category_fails_cleanly(db, blob_store):
    pipeline = _FakePipeline()
    job_id, doc_id = _seed_upload(db, blob_store, b"\x89PNG fake bytes", "image", "pic.png")

    assert _drain(db, blob_store, pipeline) is True

    job = IngestionJobRepository(db).get(job_id)
    doc = DocumentRepository(db).get(doc_id)
    assert job.status == JobStatus.FAILED
    assert "not enabled" in (job.error or "")
    assert doc.status == AssetStatus.ERROR
    assert pipeline.calls == []  # never reached indexing


def test_processor_missing_document_link_fails(db, blob_store):
    pipeline = _FakePipeline()
    # A job with no parent document (parent_id unset).
    blob_ref = blob_store.put(b"orphan")
    job_repo = IngestionJobRepository(db)
    job_id = job_repo.create(IngestionJob(
        job_type=JobType.INGEST, status=JobStatus.QUEUED, content_hash=blob_ref.content_hash,
    ))

    assert _drain(db, blob_store, pipeline) is True

    job = job_repo.get(job_id)
    assert job.status == JobStatus.FAILED
    assert "not linked" in (job.error or "").lower() or "document" in (job.error or "").lower()
    assert pipeline.calls == []
