"""
aeam/ingestion/extraction.py

Text extraction for the Enterprise Ingestion Pipeline (Phase B1.3).

The single genuinely-new capability B1.3 adds: turn the raw bytes an upload
stored in the BlobStore (Phase B1.1) into plain text that the EXISTING
IngestionPipeline (``aeam/agents/rag/ingestion_pipeline.py``) can chunk, embed,
and index. Nothing here chunks, embeds, or touches Qdrant — that is delegated
to the already-built pipeline, never re-implemented.

Scope — approved Tier 1 + Tier 2:
  - Text-native (stdlib only):   markdown, log/txt, json, xml
  - Tabular (pandas):            csv, excel
  - Documents (pure-Python):     pdf (pypdf), docx (python-docx)

Deferred to a later phase (Tier 3 — heavy/system dependencies): image (OCR),
audio / video (transcription). Those categories are accepted by the upload
validator but raise :class:`UnsupportedCategoryError` here, so their jobs fail
cleanly with a stable reason instead of silently indexing nothing.

Heavy / optional parsers (``pypdf``, ``docx``, ``openpyxl`` via pandas) are
imported LAZILY inside each extractor so a missing optional dependency degrades
only that one format rather than breaking this module's import — and so this
module stays cheap to import at application startup.

This module is pure and stateless: it performs no blob, registry, network, or
Qdrant access.
"""

from __future__ import annotations

import io
import json
import re
import html
from dataclasses import dataclass, field
from typing import Any, Callable


# ---------------------------------------------------------------------------
# Result + errors
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class ExtractionResult:
    """
    The outcome of extracting text from one uploaded file.

    Attributes:
        text:   The extracted plain text, ready to hand to the IngestionPipeline.
        detail: Small, format-specific facts (e.g. ``{"pages": 12}``,
                ``{"rows": 480, "columns": 9}``) — informational only, surfaced
                on the job/document for observability. Never load-bearing.
    """

    text: str
    detail: dict[str, Any] = field(default_factory=dict)

    @property
    def char_count(self) -> int:
        return len(self.text)


class ExtractionError(Exception):
    """
    Raised when text extraction fails.

    Mirrors :class:`~aeam.ingestion.validation.IngestValidationError`: a
    machine-stable ``reason`` code plus a human-readable ``detail``, so the
    job worker can record a structured, greppable failure.
    """

    def __init__(self, reason: str, detail: str) -> None:
        self.reason = reason
        self.detail = detail
        super().__init__(detail)


class UnsupportedCategoryError(ExtractionError):
    """
    Raised when a file's category has no registered text extractor yet.

    Distinct from a genuine extraction failure: the file may be perfectly
    valid, but its format (e.g. ``image``/``audio``/``video``) needs
    capabilities deferred to a later phase (OCR / transcription).
    """

    def __init__(self, category: str, reason: str, detail: str) -> None:
        self.category = category
        super().__init__(reason, detail)


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

# Matches an XML/HTML tag so its content can be recovered without parsing the
# document (avoids XXE and entity-expansion DoS on untrusted uploads).
_TAG_RE = re.compile(r"<[^>]+>")
_WS_RE = re.compile(r"[ \t\r\f\v]+")


def _decode_text(data: bytes) -> str:
    """
    Decode bytes to text, tolerant of the encodings enterprise files arrive in.

    Tries UTF-8 (BOM-aware) then Latin-1 (which never raises, mapping every
    byte), so a mis-labelled file degrades to best-effort text rather than
    crashing the job.
    """
    for enc in ("utf-8-sig", "utf-8"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("latin-1", errors="replace")


def _import_pandas():
    try:
        import pandas as pd  # noqa: PLC0415 (lazy by design)
        return pd
    except ImportError as exc:  # pragma: no cover - pandas is a core dependency
        raise ExtractionError(
            "missing_dependency", "pandas is required to extract tabular files."
        ) from exc


# ---------------------------------------------------------------------------
# Extractors — one per category. Signature: (bytes, filename|None) -> Result
# ---------------------------------------------------------------------------

def _extract_text_native(data: bytes, filename: str | None) -> ExtractionResult:
    """Markdown / log / plain text — decode as-is (already human text)."""
    text = _decode_text(data).strip()
    if not text:
        raise ExtractionError("empty_content", "File contained no text after decoding.")
    return ExtractionResult(text=text, detail={"format": "text"})


def _extract_json(data: bytes, filename: str | None) -> ExtractionResult:
    """JSON — validate it parses, then pretty-print so keys/values embed well."""
    raw = _decode_text(data)
    try:
        parsed = json.loads(raw)
    except (json.JSONDecodeError, ValueError) as exc:
        raise ExtractionError("invalid_json", f"File is not valid JSON: {exc}") from exc
    text = json.dumps(parsed, indent=2, ensure_ascii=False, default=str).strip()
    if not text:
        raise ExtractionError("empty_content", "JSON document was empty.")
    return ExtractionResult(text=text, detail={"format": "json"})


def _extract_xml(data: bytes, filename: str | None) -> ExtractionResult:
    """
    XML — recover the human-readable text content.

    Strips tags with a regex and unescapes character references rather than
    parsing, which deliberately avoids the XXE / billion-laughs entity-expansion
    risks of parsing untrusted XML with a stdlib parser. Tag NAMES are dropped;
    element/attribute *content* is what matters for retrieval.
    """
    raw = _decode_text(data)
    stripped = _TAG_RE.sub(" ", raw)          # remove real tags first ...
    unescaped = html.unescape(stripped)       # ... then decode &amp; etc. in content
    text = _WS_RE.sub(" ", unescaped)
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip()).strip()
    if not text:
        raise ExtractionError("empty_content", "XML file contained no extractable text.")
    return ExtractionResult(text=text, detail={"format": "xml"})


def _extract_csv(data: bytes, filename: str | None) -> ExtractionResult:
    """CSV — parse with pandas (robust quoting/encoding), re-emit as clean text."""
    pd = _import_pandas()
    try:
        df = pd.read_csv(io.BytesIO(data))
    except UnicodeDecodeError:
        df = pd.read_csv(io.BytesIO(data), encoding="latin-1")
    except Exception as exc:  # noqa: BLE001 - EmptyDataError, ParserError, ...
        raise ExtractionError("csv_parse_error", f"Could not parse CSV: {exc}") from exc
    if df.empty:
        raise ExtractionError("empty_content", "CSV file contained no data rows.")
    text = df.to_csv(index=False)
    return ExtractionResult(
        text=text, detail={"rows": int(df.shape[0]), "columns": int(df.shape[1])}
    )


def _extract_excel(data: bytes, filename: str | None) -> ExtractionResult:
    """Excel — read every sheet with pandas (openpyxl engine), concatenate as text."""
    pd = _import_pandas()
    try:
        sheets = pd.read_excel(io.BytesIO(data), sheet_name=None)  # {name: DataFrame}
    except ImportError as exc:
        raise ExtractionError(
            "missing_dependency",
            "openpyxl (for .xlsx) or xlrd (for legacy .xls) is required to read Excel files.",
        ) from exc
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError("excel_parse_error", f"Could not parse Excel workbook: {exc}") from exc

    parts: list[str] = []
    total_rows = 0
    for name, df in sheets.items():
        if df.empty:
            continue
        total_rows += int(df.shape[0])
        parts.append(f"# Sheet: {name}\n{df.to_csv(index=False)}")
    if not parts:
        raise ExtractionError("empty_content", "Excel workbook contained no data rows.")
    return ExtractionResult(
        text="\n\n".join(parts), detail={"sheets": len(sheets), "rows": total_rows}
    )


def _extract_pdf(data: bytes, filename: str | None) -> ExtractionResult:
    """PDF — extract the embedded text layer with pypdf (no OCR)."""
    try:
        import pypdf  # noqa: PLC0415 (lazy by design)
    except ImportError as exc:
        raise ExtractionError(
            "missing_dependency", "pypdf is required to extract text from PDF files."
        ) from exc
    try:
        reader = pypdf.PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:  # noqa: BLE001 - encrypted/corrupt PDFs, etc.
        raise ExtractionError("pdf_parse_error", f"Could not read PDF: {exc}") from exc

    text = "\n\n".join(p for p in pages if p.strip()).strip()
    if not text:
        raise ExtractionError(
            "no_text_layer",
            "PDF has no extractable text layer (likely a scanned document); "
            "OCR is not enabled in this phase.",
        )
    return ExtractionResult(text=text, detail={"pages": len(pages)})


def _extract_docx(data: bytes, filename: str | None) -> ExtractionResult:
    """DOCX — join paragraph text and table cell text with python-docx."""
    try:
        import docx  # noqa: PLC0415 (python-docx; lazy by design)
    except ImportError as exc:
        raise ExtractionError(
            "missing_dependency", "python-docx is required to extract text from DOCX files."
        ) from exc
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001
        raise ExtractionError("docx_parse_error", f"Could not read DOCX: {exc}") from exc

    parts: list[str] = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        raise ExtractionError("empty_content", "DOCX contained no extractable text.")
    return ExtractionResult(text=text, detail={"paragraphs": len(document.paragraphs)})


# ---------------------------------------------------------------------------
# Registry + public dispatch
# ---------------------------------------------------------------------------

#: category (from validation.SUPPORTED_EXTENSIONS values) -> extractor callable.
_EXTRACTORS: dict[str, Callable[[bytes, str | None], ExtractionResult]] = {
    "markdown": _extract_text_native,
    "log": _extract_text_native,
    "json": _extract_json,
    "xml": _extract_xml,
    "csv": _extract_csv,
    "excel": _extract_excel,
    "pdf": _extract_pdf,
    "docx": _extract_docx,
}

#: Categories with a real text extractor in this phase (Tier 1 + Tier 2).
PROCESSABLE_CATEGORIES: frozenset[str] = frozenset(_EXTRACTORS)

#: Categories accepted at upload but whose extraction is deferred (Tier 3).
DEFERRED_CATEGORIES: frozenset[str] = frozenset({"image", "audio", "video"})


def can_extract(category: str) -> bool:
    """Return ``True`` if this phase can extract text from ``category``."""
    return category in _EXTRACTORS


def extract_text(
    data: bytes,
    category: str,
    filename: str | None = None,
) -> ExtractionResult:
    """
    Extract plain text from an uploaded file's bytes.

    Args:
        data:     The original file bytes (as read back from the BlobStore).
        category: The format category resolved at upload
                  (see :data:`aeam.ingestion.validation.SUPPORTED_EXTENSIONS`).
        filename: Original filename, for logging/heuristics only (optional).

    Returns:
        An :class:`ExtractionResult` with non-empty ``text``.

    Raises:
        UnsupportedCategoryError: ``category`` has no extractor yet (deferred
                                  Tier-3 format, or an unknown category).
        ExtractionError:          Extraction was attempted but failed (parse
                                  error, empty content, missing optional
                                  dependency, scanned PDF, ...).
    """
    if data is None or len(data) == 0:
        raise ExtractionError("empty_content", "No data provided to extract.")

    extractor = _EXTRACTORS.get(category)
    if extractor is None:
        if category in DEFERRED_CATEGORIES:
            raise UnsupportedCategoryError(
                category,
                "deferred_category",
                f"Text extraction for '{category}' files is not enabled yet "
                f"(Tier 3 — OCR / transcription, deferred to a later phase).",
            )
        raise UnsupportedCategoryError(
            category,
            "unknown_category",
            f"No text extractor is registered for category '{category}'.",
        )

    result = extractor(data, filename)
    if not result.text or not result.text.strip():
        # Defensive: every extractor already guards this, but never let empty
        # text reach the indexing pipeline.
        raise ExtractionError("empty_content", f"Extraction produced no text for a '{category}' file.")
    return result
